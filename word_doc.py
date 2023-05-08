from docx import Document
import os
word = Document()
word.add_heading("My_name is manlow",level=2)
word.add_paragraph("This is my paragraph")
if os.path.exists(os.path.join(os.getcwd(),"Manlow.docx")):
    os.remove("Manlow.docx")
word.save("Manlow.docx")
def page_break(doc):
    doc.add_page_break()

def add_paragraph(doc,paragraph):
    doc.add_pagraph(paragraph)

def add_picture(doc,path):
    doc.add_picture(path)